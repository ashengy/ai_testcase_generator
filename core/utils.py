# core/utils.py
import json
import os
import re
import traceback
import xml.etree.cElementTree as ET

import cv2
import numpy as np
import pandas as pd
import yaml
from PyPDF2 import PdfReader
from docx import Document
from paddleocr import PaddleOCR

from config.constants import TEMPLATE_PHRASES, CONTENT_FILTER_FUZZY, CONTENT_FILTER_EXACT


def clean_headers_footers(content):
    """清理页眉、页脚和目录内容"""
    cleaned_content = {"paragraphs": [], "tables": content["tables"]}

    headers_footers = content.get("headers", []) + content.get("footers", [])
    for paragraph in content["paragraphs"]:
        if paragraph in headers_footers:
            continue
        if re.match(r"^\s*(第[\d一二三四五六七八九十]+章|\d+(\.\d+)*).*$", paragraph):
            continue
        if re.match(r"\\t", paragraph):
            continue
        cleaned_content["paragraphs"].append(paragraph)

    return cleaned_content


def remove_template_phrases(content):
    """删除文档中的模板固有内容"""
    cleaned_paragraphs = []
    for paragraph in content["paragraphs"]:
        if any(phrase in paragraph for phrase in TEMPLATE_PHRASES):
            continue
        cleaned_paragraphs.append(paragraph)
    return {"paragraphs": cleaned_paragraphs, "tables": content["tables"]}


def clean_text(text):
    """文本清洗处理"""
    text = clean_headers_footers(text)
    return text


def remove_toc(doc):
    """去除 Word 文档中的目录"""
    paragraphs_to_remove = []

    for paragraph in doc.paragraphs:
        if paragraph.style is not None and "TOC" in paragraph.style.name:
            paragraphs_to_remove.append(paragraph)
        elif paragraph.text.strip() and (paragraph.text[0].isdigit() or paragraph.text.startswith("1.")):
            paragraphs_to_remove.append(paragraph)

    for paragraph in paragraphs_to_remove:
        p = paragraph._element
        p.getparent().remove(p)
    return doc


def is_heading_enhanced(paragraph):
    """判断段落是否为标题"""
    heading_features = {
        'keywords': ['标题', 'heading', 'header', 'h1', 'h2', 'h3', 'chapter'],
        'font_size': (14, 72),
        'bold_threshold': 0.7
    }

    style_match = any(
        kw in (paragraph.style.name or "").lower()
        for kw in heading_features['keywords']
    )
    if style_match:
        return True

    try:
        font = paragraph.style.font or paragraph.document.styles['Normal'].font
        effective_size = font.size.pt if font.size else 12
        is_bold = font.bold or False
        size_ok = heading_features['font_size'][0] <= effective_size <= heading_features['font_size'][1]
        bold_ok = is_bold >= heading_features['bold_threshold']
        return size_ok and bold_ok
    except Exception as e:
        print(f"格式检查失败: {str(e)}")
        return False


def get_target_pic(file, target_title):
    """提取目标标题下的图片"""
    namespace = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'v': "urn:schemas-microsoft-com:vml",
        'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
        'pic': "http://schemas.openxmlformats.org/drawingml/2006/picture"
    }

    doc = Document(file)
    target_title_list = None
    if target_title:
        target_title_list = target_title.strip(' ').split(',')

    def get_img(root_element, target_tag, target_attribute, out_list):
        for child in root_element:
            tag = child.tag
            attribute = child.attrib
            if tag in target_tag and target_attribute in child.attrib.keys():
                target_value = child.attrib[target_attribute]
                out_list.append(target_value)
            else:
                get_img(child, target_tag, target_attribute, out_list)

    xml_element = []
    text_content = []
    found_start = False
    if target_title_list:
        n = 0
        found_times = len(target_title_list)
        for par in doc.paragraphs:
            for title in target_title_list:
                if title in par.text and 'toc' not in par.style.name.lower() and is_heading_enhanced(par):
                    found_start = True
                    n += 1
                    break
                elif title not in par.text and is_heading_enhanced(par):
                    found_start = False
            if found_start and n <= found_times:
                if par.text == '':
                    pass
                text_content.append(par.text if par.text else '+')
                xml_element.append(par._element.xml)

    rId = []
    id = []
    for element in xml_element:
        if element:
            root = ET.fromstring(element)
            target_tag = [
                '{urn:schemas-microsoft-com:vml}imagedata',
                '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
            ]
            target_attribute1 = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
            target_attribute2 = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'

            get_img(root, target_tag, target_attribute1, rId)
            get_img(root, target_tag, target_attribute2, id)

    j = 0
    dd = id + rId
    images = []
    for id_val in dd:
        img_part = doc.part.related_parts[id_val]
        img_binary = img_part.blob
        img = cv2.imdecode(np.frombuffer(img_binary, np.uint8), cv2.IMREAD_COLOR)
        try:
            if not img.any():
                j += 1
                continue
            else:
                img_name = os.path.dirname(file) + r"\img" + str(j) + ".jpg"
                cv2.imwrite(img_name, img)
                images.append(img_name)
                j += 1
        except ValueError as e:
            print(e)
            img_name = os.path.dirname(file) + r"\img" + str(j) + ".jpg"
            cv2.imwrite(img_name, img)
            images.append(img_name)
            j += 1

    return images if images else None


def perform_ocr_with_paddle(images):
    """使用 PaddleOCR 对图片进行文字识别"""
    results = []
    ocr = PaddleOCR(use_angle_cls=True, lang="ch")
    for image_path in images:
        try:
            img = cv2.imread(image_path)
            ocr_result = ocr.ocr(img, cls=True)
            text_lines = [line[1][0] for line in ocr_result[0]]
            results.append((image_path, "\n".join(text_lines)))
        except Exception as e:
            results.append((image_path, f"OCR 识别失败: {e}"))
    return results


def extract_text_by_title(docx_path, title_keywords, table_keywords, pic_keywords):
    """提取多个标题下的正文内容"""
    print(f"需求文档路径:{docx_path}")
    if not os.path.exists(docx_path):
        print("文件不存在")
        return {}

    try:
        doc = Document(docx_path)
        print("文件加载成功")
    except Exception as e:
        print(f"加载文件时发生错误: {e}")
        return {}

    result = {}
    doc = remove_toc(doc)
    for title_keyword in title_keywords.split(','):
        content = []
        capture = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                if title_keyword in text and (paragraph.style is None or 'toc' not in paragraph.style.name.lower()):
                    capture = True
                    content.append(text)
                    continue
                if capture and re.match(r"^\d+(\.\d+)*\s+.+", text):
                    break
                elif paragraph.style is not None and "标题" in paragraph.style.name and capture:
                    break
                if capture and text:
                    content.append(text)
        result[title_keyword] = "\n".join(content)

    pic_keyword_list = []
    if pic_keywords:
        pic_keyword_list = pic_keywords.strip(' ').split(',')
    if len(pic_keyword_list) == 0:
        print("未输入图片标题，不进行图片信息提取及OCR识别")

    for pic_keyword in pic_keyword_list:
        image_paths = get_target_pic(docx_path, pic_keyword)
        if not image_paths:
            print("未找到目标标题下的图片")
        else:
            print(f"提取到 {len(image_paths)} 张图片：{image_paths}")
            print("\n正在进行 OCR 识别...")
            ocr_results = perform_ocr_with_paddle(image_paths)
            ocr_results_text = ''
            for image_path, text in ocr_results:
                ocr_results_text += text + '--'
            result['识别内容'] = ocr_results_text

    return result


def is_title(paragraph, filter_list=None):
    """判断段落是否为标题"""
    text = paragraph.text.strip()
    if not text:
        return False

    if paragraph.style.name.startswith("Heading"):
        return True

    if re.match(r"^\d+(\.\d+)*\s+.*", text):
        return True
    if re.match(r"^\d+[-]\d+\s+.*", text):
        return True
    if re.match(r"^(附录|参考文献|功能要求|概述).*$", text):
        return True

    if filter_list and any(keyword in text for keyword in filter_list):
        return True

    return False


def extract_content(file_path, image_folder="extracted_images"):
    """提取需求文档内容，包括段落、表格，并处理图片"""
    doc = Document(file_path)
    doc = remove_toc(doc)
    content = {"paragraphs": [], "tables": [], "images": []}

    if not os.path.exists(image_folder):
        os.makedirs(image_folder)

    for paragraph in doc.paragraphs:
        skip_section = False
        text = paragraph.text.strip()

        if is_title(paragraph, TEMPLATE_PHRASES):
            if any(keyword in text for keyword in TEMPLATE_PHRASES) or 'toc' in paragraph.style.name.lower():
                skip_section = True
                continue
            else:
                skip_section = False
        else:
            for ele in CONTENT_FILTER_FUZZY:
                if ele in text:
                    skip_section = True
            for ele in CONTENT_FILTER_EXACT:
                if ele == text:
                    skip_section = True

            if not skip_section and text:
                content["paragraphs"].append(text)

    print(f'----------------\n{content}')

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append(table_data)

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_name = os.path.join(image_folder, os.path.basename(rel.target_ref))
            with open(image_name, "wb") as img_file:
                img_file.write(image_data)
            content["images"].append(image_name)

    return content


def read_file(file_path):
    """多格式文件读取"""
    try:
        if file_path.endswith('.docx'):
            try:
                return extract_content(file_path)
            except Exception as e:
                print(f"读取docx文件出错: {e}")
                traceback.print_exc()
                return ""
        elif file_path.endswith('.pdf'):
            try:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    pages_text = []
                    for page in reader.pages:
                        try:
                            text = page.extract_text()
                            if text:
                                text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
                                text = text.encode('utf-8', errors='ignore').decode('utf-8')
                            pages_text.append(text if text else "")
                        except Exception as page_e:
                            print(f"提取PDF页面文本时出错: {page_e}")
                            pages_text.append("")
                    combined_text = '\n'.join(pages_text)
                    combined_text = '\n'.join(line for line in combined_text.split('\n') if line.strip())
                    return {"paragraphs": [combined_text], "tables": []}
            except Exception as e:
                print(f"读取pdf文件出错: {e}")
                traceback.print_exc()
                return {"paragraphs": [""], "tables": []}
        elif file_path.endswith('.xlsx'):
            try:
                df = pd.read_excel(file_path)
                return df.to_markdown(index=False)
            except Exception as e:
                error_msg = f"无法读取Excel文件: {str(e)}"
                print(error_msg)
                traceback.print_exc()
                return ""
        elif file_path.endswith('.md'):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                print(f"读取markdown文件出错: {e}")
                traceback.print_exc()
                return ""
        elif file_path.endswith('.txt'):
            try:
                encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError:
                        continue
                raise Exception("无法使用常见编码读取文件")
            except Exception as e:
                print(f"读取txt文件出错: {e}")
                traceback.print_exc()
                return ""
        elif file_path.endswith('.json'):
            try:
                encodings = ['utf-8', 'gbk', 'gb2312']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            return json.load(f)
                    except UnicodeDecodeError:
                        continue
                raise Exception("无法使用常见编码读取JSON文件")
            except Exception as e:
                print(f"读取json文件出错: {e}")
                traceback.print_exc()
                return ""
        elif file_path.endswith(('.yaml', '.yml')):
            try:
                encodings = ['utf-8', 'gbk', 'gb2312']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            return yaml.safe_load(f)
                    except UnicodeDecodeError:
                        continue
                raise Exception("无法使用常见编码读取YAML文件")
            except Exception as e:
                print(f"读取yaml文件出错: {e}")
                traceback.print_exc()
                return ""
    except Exception as e:
        error_msg = f"读取文件时发生未知错误: {str(e)}"
        print(error_msg)
        traceback.print_exc()
        return ""


def chunk_xlsx(df, max_chunk_size=1000):
    """将 .xlsx 接口文档按接口定义分块，确保接口数据完整"""
    print('开始对xlsx文档内容进行分块！')
    chunks = []
    current_chunk = ""
    current_size = 0

    for _, row in df.iterrows():
        row_content = row.to_json(force_ascii=False)
        row_size = len(row_content)

        if current_size + row_size > max_chunk_size:
            chunks.append(current_chunk)
            current_chunk = row_content
            current_size = row_size
        else:
            current_chunk += "\n" + row_content
            current_size += row_size
    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def chunk_yaml(data, max_chunk_size=1000):
    """将 .yml 接口文档按接口定义分块，确保接口数据完整"""
    chunks = []
    current_chunk = ""
    current_size = 0
    print("开始对yml文件进行分块~~")

    if isinstance(data, dict):
        for key, value in data.items():
            interface_content = yaml.dump({key: value})
            interface_size = len(interface_content)

            if current_size + interface_size > max_chunk_size:
                chunks.append(current_chunk)
                current_chunk = interface_content
                current_size = interface_size
            else:
                current_chunk += "\n" + interface_content
                current_size += interface_size
    elif isinstance(data, list):
        for data_ in data:
            for key, value in data_.items():
                interface_content = yaml.dump({key: value}, allow_unicode=True)
                interface_size = len(interface_content)

                if current_size + interface_size > max_chunk_size:
                    chunks.append(current_chunk)
                    current_chunk = interface_content
                    current_size = interface_size
                else:
                    current_chunk += "\n" + interface_content
                    current_size += interface_size

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def chunk_json(content, max_chunk_size=1000):
    """将 .json 接口文档按接口定义分块，确保接口数据完整"""
    print("开始对json或yml文件进行分块~~")
    chunks = []
    current_chunk = ""
    current_size = 0

    for n, interface in enumerate(content):
        interface_content = json.dumps(interface, indent=2, ensure_ascii=False)
        interface_size = len(interface_content)

        if current_size + interface_size > max_chunk_size:
            if n == 0:
                chunks.append(interface_content)
            else:
                chunks.append(current_chunk)
            current_chunk = interface_content
            current_size = interface_size
            n += 1
        else:
            current_chunk += "\n" + interface_content
            current_size += interface_size
            n += 1

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def json_to_excel(json_data, output_file):
    """将任意 JSON 数据中的键作为表头，值作为值，转换为 Excel 表格"""
    data_list = None

    if isinstance(json_data, str):
        json_data = json.loads(json_data)

    if isinstance(json_data, dict):
        for key, value in json_data.items():
            if isinstance(value, list):
                data_list = value
                break
    elif isinstance(json_data, list):
        data_list = json_data
    else:
        raise ValueError("JSON 数据中未找到列表部分！")

    if data_list:
        df = pd.DataFrame(data_list)
        df.to_excel(output_file, index=False)
        print(f"Excel 文件已成功生成：{output_file}")
    else:
        print("Excel 文件生成失败")

def chunk_text(text, chunk_size=2000, overlap=300):
    """
    将文本按固定长度分块，同时添加滑动窗口重叠。

    参数：
    - text (str): 输入的长文本内容
    - chunk_size (int): 每块的最大字符数
    - overlap (int): 相邻块的重叠字符数

    返回：
    - list: 分块后的文本列表
    """
    # print("开始对文本进行分块：",text)
    if text:
        clear_text = text.replace("·","") #去掉文档里的多余符号
        chunks = []
        start = 0
        text_length = len(clear_text)

        while start < text_length:
            end = min(start + chunk_size, text_length)
            chunk = clear_text[start:end]
            chunks.append(chunk)
            # 滑动窗口：下一块的起始位置向后移动 chunk_size - overlap
            start += chunk_size - overlap
        print(f"分块完成，共生成了 {len(chunks)} 个块。")
        return chunks
