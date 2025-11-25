import io
from typing import List, Tuple, Dict

from PIL import Image
from docx import Document
from docx.oxml.ns import qn
from lxml import etree


def insert_image_position_with_list(
        doc_path: str,
        image_replacement_list: List[str],
        min_width: int = 100,
        min_height: int = 100
) -> str:
    """
    完美保留文档结构：正文与表格穿插出现，表格内图片替换后仍在原位置
    核心逻辑：按文档实际块级元素顺序（段落→表格→段落...）处理，不改变原有布局
    """
    doc = Document(doc_path)
    namespaces = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    # -------------------------- 关键步骤1：解析文档所有块级元素（按实际顺序）--------------------------
    # 块级元素类型：'paragraph'（正文段落）或 'table'（表格）
    block_elements = []  # 存储：(block_type, index, element)，index是该类型内的索引
    para_index = 0
    table_index = 0

    # 遍历文档的body元素，按XML顺序提取块级元素（最精准的文档顺序）
    body = doc.element.body
    for child in body:
        tag = child.tag.split('}')[-1]  # 获取标签名（去除命名空间）
        if tag == 'p':  # 段落
            block_elements.append(('paragraph', para_index, doc.paragraphs[para_index]))
            para_index += 1
        elif tag == 'tbl':  # 表格
            block_elements.append(('table', table_index, doc.tables[table_index]))
            table_index += 1

    # for idx, (block_type, inner_idx, _) in enumerate(block_elements):
    #     print(f"   块{idx + 1}：{block_type}（内部索引{inner_idx + 1}）")

    # -------------------------- 关键步骤2：收集所有图片（关联到块级元素顺序）--------------------------
    image_meta = []  # 存储：(block_order, block_type, block_inner_idx, position_info, para_key, run_idx, blip_idx, r_id, width, height)
    r_id_set = set()
    total_images = 0
    global_image_idx = 0

    for block_order, (block_type, block_inner_idx, block_elem) in enumerate(block_elements):
        if block_type == 'paragraph':
            # 处理正文段落
            paragraph = block_elem
            for run_idx, run in enumerate(paragraph.runs):
                try:
                    run_xml = etree.fromstring(etree.tostring(run.element))
                    blip_elements = run_xml.xpath('.//a:blip', namespaces=namespaces)
                    for blip_idx, blip in enumerate(blip_elements):
                        r_id = blip.get(qn('r:embed'))
                        if not (r_id and r_id not in r_id_set and r_id in doc.part.related_parts):
                            continue

                        # 图片尺寸过滤
                        part = doc.part.related_parts[r_id]
                        img = Image.open(io.BytesIO(part.blob))
                        width, height = img.size
                        if width < min_width or height < min_height:
                            print(
                                f"跳过极小图：块{block_order + 1}（正文段落{block_inner_idx + 1}）→Run{run_idx + 1}→第{blip_idx + 1}张 | {width}x{height} | r_id:{r_id}")
                            continue

                        # 位置信息
                        position_info = f"正文段落{block_inner_idx + 1}→Run{run_idx + 1}→第{blip_idx + 1}张图"
                        para_key = f"para_{block_inner_idx}"  # 正文段落唯一标识
                        # 存储元信息（block_order决定整体顺序）
                        image_meta.append((
                            block_order, block_type, block_inner_idx, position_info,
                            para_key, run_idx, blip_idx, r_id, width, height
                        ))
                        r_id_set.add(r_id)
                        total_images += 1
                        global_image_idx += 1
                        print(
                            f"图片{global_image_idx}：块{block_order + 1}→{position_info} | {width}x{height} | r_id:{r_id}")
                except Exception as e:
                    print(
                        f"收集正文图片错误：块{block_order + 1}（正文段落{block_inner_idx + 1}）→Run{run_idx + 1} | 错误：{str(e)}")
                    continue

        elif block_type == 'table':
            # 处理表格（保留表格内位置）
            table = block_elem
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for cell_para_idx, paragraph in enumerate(cell.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            try:
                                run_xml = etree.fromstring(etree.tostring(run.element))
                                blip_elements = run_xml.xpath('.//a:blip', namespaces=namespaces)
                                for blip_idx, blip in enumerate(blip_elements):
                                    r_id = blip.get(qn('r:embed'))
                                    if not (r_id and r_id not in r_id_set and r_id in doc.part.related_parts):
                                        continue

                                    # 图片尺寸过滤
                                    part = doc.part.related_parts[r_id]
                                    img = Image.open(io.BytesIO(part.blob))
                                    width, height = img.size
                                    if width < min_width or height < min_height:
                                        print(
                                            f"跳过极小图：块{block_order + 1}（表格{block_inner_idx + 1}）→行{row_idx + 1}→列{cell_idx + 1}→段落{cell_para_idx + 1}→Run{run_idx + 1} | {width}x{height} | r_id:{r_id}")
                                        continue

                                    # 位置信息
                                    position_info = f"表格{block_inner_idx + 1}→行{row_idx + 1}→列{cell_idx + 1}→段落{cell_para_idx + 1}→Run{run_idx + 1}→第{blip_idx + 1}张图"
                                    para_key = f"table_{block_inner_idx}_cell_{cell_idx}_para_{cell_para_idx}"  # 表格内段落唯一标识
                                    # 存储元信息（block_order决定表格在文档中的位置）
                                    image_meta.append((
                                        block_order, block_type, block_inner_idx, position_info,
                                        para_key, run_idx, blip_idx, r_id, width, height
                                    ))
                                    r_id_set.add(r_id)
                                    total_images += 1
                                    global_image_idx += 1
                                    print(
                                        f"    图片{global_image_idx}：块{block_order + 1}→{position_info} | {width}x{height} | r_id:{r_id}")
                            except Exception as e:
                                print(
                                    f"    收集表格图片错误：块{block_order + 1}（表格{block_inner_idx + 1}）→行{row_idx + 1}→列{cell_idx + 1}→段落{cell_para_idx + 1}→Run{run_idx + 1} | 错误：{str(e)}")
                                continue

    print(f"图片收集完成：共识别到 {total_images} 张有效图片")
    if total_images == 0:
        print("未识别到任何有效图片，返回原文档文本")
        return get_document_with_original_structure(doc, block_elements)

    # -------------------------- 按文档实际顺序排序图片（关键）--------------------------
    # 排序键：block_order（块级元素顺序）→ 内部位置，确保与文档阅读顺序完全一致
    image_meta.sort(key=lambda x: (x[0], x[5], x[6]))  # block_order → run_idx → blip_idx
    for replace_idx, (block_order, block_type, block_inner_idx, position_info, _, _, _, _, width, height) in enumerate(
            image_meta):
        replace_text = image_replacement_list[replace_idx] if replace_idx < len(image_replacement_list) else ""

    # -------------------------- 处理替换列表长度 --------------------------
    if len(image_replacement_list) < total_images:
        fill_num = total_images - len(image_replacement_list)
        image_replacement_list.extend([""] * fill_num)
        print(
            f"\n替换列表长度不足（{len(image_replacement_list) - fill_num} < {total_images}），已补充 {fill_num} 个空字符串")
    elif len(image_replacement_list) > total_images:
        image_replacement_list = image_replacement_list[:total_images]
        print(
            f"\n⚠替换列表长度超出（{len(image_replacement_list) + (len(image_replacement_list) - total_images)} > {total_images}），已截取前 {total_images} 个元素")

    # 建立替换映射：(para_key, run_idx, blip_idx) → 替换文字
    replace_map: Dict[Tuple[str, int, int], str] = {}
    for replace_idx, (_, _, _, _, para_key, run_idx, blip_idx, _, _, _) in enumerate(image_meta):
        replace_map[(para_key, run_idx, blip_idx)] = image_replacement_list[replace_idx]

    # -------------------------- 按文档原结构执行替换（关键）--------------------------

    final_content_parts = []

    for block_order, (block_type, block_inner_idx, block_elem) in enumerate(block_elements):
        if block_type == 'paragraph':
            # 处理正文段落
            paragraph = block_elem
            para_content = process_single_paragraph(
                paragraph=paragraph,
                para_key=f"para_{block_inner_idx}",
                replace_map=replace_map,
                namespaces=namespaces,
                position_label=f"正文段落{block_inner_idx + 1}"
            )
            if para_content.strip():
                final_content_parts.append(para_content)

        elif block_type == 'table':
            # 处理表格（保留原表格结构和位置）
            table = block_elem
            table_content = process_table(
                table=table,
                table_index=block_inner_idx,
                replace_map=replace_map,
                namespaces=namespaces
            )
            final_content_parts.append(table_content)

    # 合并内容，用两个换行符分隔块级元素（保持原文档间距）
    final_content = "\n".join(final_content_parts)
    print(f"\n替换完成！表格已保留在原位置，图片替换准确")
    return final_content


def process_single_paragraph(
        paragraph,
        para_key: str,
        replace_map: Dict[Tuple[str, int, int], str],
        namespaces: Dict[str, str],
        position_label: str
) -> str:
    """处理单个段落（正文或表格单元格内段落）的图片替换"""
    para_content = ""
    for run_idx, run in enumerate(paragraph.runs):
        run_text = run.text or ""
        try:
            run_xml = etree.fromstring(etree.tostring(run.element))
            blip_elements = run_xml.xpath('.//a:blip', namespaces=namespaces)
        except Exception as e:
            print(f"处理Run错误：{position_label}→Run{run_idx + 1} | 错误：{str(e)}")
            blip_elements = []

        if blip_elements:
            # 有图片：先加文本，再按顺序替换图片
            if run_text:
                para_content += run_text
            for blip_idx, blip in enumerate(blip_elements):
                key = (para_key, run_idx, blip_idx)
                if key in replace_map:
                    replace_text = replace_map[key]
                    if replace_text:
                        para_content += f"[图片：{replace_text}]"
                    else:
                        para_content += ""  # 无效图片，拼接空字符串
                    print(
                        f"替换成功：{position_label}→Run{run_idx + 1}→第{blip_idx + 1}张图 → [图片:{replace_text}]")
                else:
                    r_id = blip.get(qn('r:embed'))
                    print(
                        f"未替换：{position_label}→Run{run_idx + 1}→第{blip_idx + 1}张图（r_id:{r_id}，无匹配替换规则）")
        else:
            # 无图片：直接加文本
            if run_text:
                para_content += run_text
    return para_content


def process_table(
        table,
        table_index: int,
        replace_map: Dict[Tuple[str, int, int], str],
        namespaces: Dict[str, str]
) -> str:
    """处理表格，替换单元格内图片，保留表格结构"""
    table_content = []
    for row_idx, row in enumerate(table.rows):
        row_cells = []
        for cell_idx, cell in enumerate(row.cells):
            cell_para_contents = []
            for cell_para_idx, paragraph in enumerate(cell.paragraphs):
                # 表格内段落的唯一标识
                para_key = f"table_{table_index}_cell_{cell_idx}_para_{cell_para_idx}"
                # 处理单元格内的单个段落
                para_content = process_single_paragraph(
                    paragraph=paragraph,
                    para_key=para_key,
                    replace_map=replace_map,
                    namespaces=namespaces,
                    position_label=f"表格{table_index + 1}→行{row_idx + 1}→列{cell_idx + 1}→段落{cell_para_idx + 1}"
                )
                if para_content.strip():
                    cell_para_contents.append(para_content)
            # 单元格内段落用换行分隔
            cell_content = "\n".join(cell_para_contents)
            row_cells.append(cell_content)
        # 表格行用制表符分隔（模拟表格列）
        table_content.append("\t".join(row_cells))
    # 表格行用换行分隔
    return "\n".join(table_content)


def get_document_with_original_structure(doc: Document, block_elements) -> str:
    """获取原文档结构的文本（无图片替换时使用）"""
    content_parts = []
    for block_type, block_inner_idx, block_elem in block_elements:
        if block_type == 'paragraph':
            para_text = block_elem.text.strip()
            if para_text:
                content_parts.append(para_text)
        elif block_type == 'table':
            table_content = []
            for row in block_elem.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    row_cells.append(cell_text)
                table_content.append("\t".join(row_cells))
            content_parts.append("\n".join(table_content))
    return "\n\n".join(content_parts)


# 调用示例
if __name__ == "__main__":
    # 替换列表：按调试日志中的「替换顺序」编写（与文档实际图片顺序完全一致）
    replacements = [
        "图1：好友系统架构图",
        "图2：联系人添加流程图",
        "",  # 表格内的图片（现在会在原表格位置）
        " ",
        "图5：数据存储结构图",
        "图6：表格内功能对比表",  # 另一张表格内的图片
        "图7：权限控制示意图",
        "图8：异常处理流程图"
    ]

    content = insert_image_position_with_list(
        doc_path=r"D:\Download\好友系统联系人系统屏蔽系统.docx",
        image_replacement_list=replacements
    )
    print(content)
