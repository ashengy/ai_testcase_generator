# ui_test/main.py
import json
import os
import sys

from PyQt5 import QtGui
from PyQt5.QtCore import Qt, QSettings
from PyQt5.QtGui import QTextCursor
from PyQt5.QtWidgets import (QMainWindow, QAbstractItemView,
                             QLineEdit, QFileDialog, QMessageBox, QListWidgetItem,
                             QApplication, QFrame, QGraphicsDropShadowEffect, QDialog)

import config.constants
from config.constants import TEMPLATE_PHRASES, CONTENT_FILTER_FUZZY, CONTENT_FILTER_EXACT, CLEAN_FLAG, design_methods
from core.pdf_image_replace import extract_pdf_text_with_image_list
from core.utils import chunk_text,normalize_data
from core.worker import GenerateThread, ImageAnalyzer
from ui.ui_deepseektool import Ui_DeepSeekTool
from ui.ui_style import load_stylesheet
from core.word_image_replace import insert_image_position_with_list

class DeepSeekTool(QMainWindow, Ui_DeepSeekTool):
    def __init__(self):
        super().__init__()
        self.setWindowIcon(self.load_icon("favicon.ico"))
        # 初始化设置类
        self.settings = QSettings("soc", "Ai-Testcase")
        self.context = None
        self.context_chunks = []
        # 预解析缓存：避免重复解析同一个 docx/pdf 导致预览卡顿
        # key: 文件绝对路径, value: 预览区应显示/用于后续生成的文本
        self._context_cache = {}
        # 记录最近一次执行过 AI 图片分析的文件，方便把分析结果写回缓存
        self._last_analyzed_path = None
        # Ai图片分析期间是否允许用户点击“开始推理”
        self._generate_was_enabled = True
        self.func_type = None
        self.module_input_pic = None
        self.api_key = ""  # 添加API Key属性
        self.image_api_key = ""
        self.init_ui()
        self.knowledge_bases = []
        self.current_dir = ""
        self.job_area = None
        self.setStyleSheet(load_stylesheet())
        # 使用从config导入的常量
        self.template_phrases = TEMPLATE_PHRASES
        self.content_filter_fuzzy = CONTENT_FILTER_FUZZY
        self.content_filter_exact = CONTENT_FILTER_EXACT
        self.clean_flag = CLEAN_FLAG

        # 设置多选下拉选项为只读
        self.comboBox_design_method.setEditable(True)
        self.comboBox_design_method.lineEdit().setReadOnly(True)
        self.comboBox_design_method.clearEditText()
        # 初始化自定义多选下拉选项模型
        self.custom_model = QtGui.QStandardItemModel()
        self.create_custom_model()
        self.custom_model.itemChanged.connect(self.reset_combox_text)

    def init_ui(self):
        """ 初始化界面 """
        self.setupUi(self)
        self.setGeometry(300, 200, 1400, 900)
        self.setMinimumSize(1280, 820)
        self.setWindowTitle("AI 测试用例生成工作台")
        # 按钮别名：如果新UI中已有这些属性则不需要设置
        if not hasattr(self, 'generate_btn'):
            self.generate_btn = getattr(self, 'generateButton', None)
        if not hasattr(self, 'refresh_prompt_btn'):
            self.refresh_prompt_btn = getattr(self, 'refreshPromptButton', None)

        # 填充默认api key
        self.api_key_input.setText(config.constants.DEEPSEEK_API_KEY)
        self.lineEdit_image_api_key.setText(config.constants.IMAGE_API_KEY)

        # 隐藏无用控件
        self.label_module_input.hide()
        self.module_input.hide()
        self.label_module_input_table.hide()
        self.module_input_table.hide()
        self.label_module_input_pic.hide()
        self.module_input_pic.hide()
        self.label.hide()

        self.param_choice_combo.setCurrentIndex(0)
        self.func_choice_combo.setCurrentIndex(0)
        self.comboBox.currentTextChanged.connect(self.updateLabel)
        self.comboBox.setCurrentIndex(6) # 设置行业默认选中游戏开发
        self.api_key_input.setEchoMode(QLineEdit.Password)
        self.preview_area.setReadOnly(False)
        self.result_area.setReadOnly(False)

        self.plainTextEdit_update_talking.setReadOnly(True)
        self.plainTextEdit_update_talking.setEnabled(True)
        # 设置日志框的最大字数，只保持最新消息
        self.plainTextEdit_update_talking.document().setMaximumBlockCount(10000)
        self.plainTextEdit_update_talking.moveCursor(QTextCursor.End)  # 滚动到底部

        # 窗口初始化时加载设置保存路径
        self.load_saved_paths()
        self._beautify_ui()
        self.prompt_input.setText("Role: 测试用例设计专家\n\n"
                                  "Rules:\n\n"
                                  "设计目标：\n"
                                  "通过正交分析法实现：\n"
                                  "使用正交表生成参数组合，覆盖所有参数对的交互组合\n\n"
                                  "用例数量：\n"
                                  "尽可能多（不少于15条）\n"
                                  "需求分析指南：\n"
                                  "1. 识别功能边界（系统做什么/不做什么）\n"
                                  "2. 提取业务规则（计算规则、验证规则）\n"
                                  "3. 定义用户角色及其权限\n"
                                  "4. 梳理关键业务流程（正常流、备选流、异常流）\n"
                                  "5. 标记敏感操作（审计日志、权限校验点）\n\n"
                                  "输出要求：\n"
                                  "1. 格式：结构化JSON,必须严格遵守JSON语法规范\n"
                                  "2. 不要使用JavaScript语法（如.repeat()方法）\n"
                                  "3. 对于需要重复字符的情况，请直接写出完整字符串，例如：\"aaaaaaa...\"而不是\"a\".repeat(7)"
                                  "4. 字符串长度限制测试时，请使用描述性文字如\"256个字符的a\"，而不是实际生成256个字符"
                                  "5. 字段：\n"
                                  "   - 用例编号：<模块缩写>-<3位序号>\n"
                                  "   - 用例标题：<测试目标> [正例/反例]\n"
                                  "   - 前置条件：初始化状态描述\n"
                                  "   - 操作步骤：带编号的明确步骤\n"
                                  "   - 预期结果：可验证的断言\n"
                                  "   - 优先级：P0(冒烟)/P1(核心)/P2(次要)\n"
                                  "3. 示例：\n"
                                  "[\n"
                                  "    {\n"
                                  "        \"用例编号\": \"PAY-001\",\n"
                                  "        \"用例标题\": \"支付功能 [正例]\",\n"
                                  "        \"前置条件\": \"用户已登录，购物车内已有商品\",\n"
                                  "        \"操作步骤\": [\n"
                                  "            \"1. 打开购物车页面\",\n"
                                  "            \"2. 点击结算按钮\",\n"
                                  "            \"3. 选择支付方式为支付宝支付\",\n"
                                  "            \"4. 确认支付金额为100-1000元人民币\",\n"
                                  "            \"5. 点击支付按钮\"\n"
                                  "        ],\n"
                                  "        \"预期结果\": \"支付成功，页面显示支付完成信息，余额扣减正确\",\n"
                                  "        \"优先级\": \"P1\"\n"
                                  "    }\n"
                                  "]\n\n"
                                  "质量标准：\n"
                                  "- 参数对组合覆盖率 ≥95%\n"
                                  "- 正向场景用例占比60%\n"
                                  "- 异常场景用例占比30%\n"
                                  "- 边界场景用例占比10%\n\n"
                                  "生成步骤：\n"
                                  "1. 参数建模 → 2. 场景分析 → 3. 用例生成 → 4. 交叉校验")

        self.file_list.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.file_list.setSortingEnabled(True)

        # 连接信号和槽
        self._connect_signals()

        # 添加按钮悬停时的tips
        self.refresh_prompt_btn.setToolTip("修改功能模式、设计方法后，需更新提示词")

        # 添加combboBox的悬停tips
        self.combo_kb.setToolTip("目前已支持docx和pdf")

    def _beautify_ui(self):
        """通过运行时设置优化视觉层次，不改动核心交互逻辑。"""
        self.centralwidget.setObjectName("appSurface")
        self.statusbar.setSizeGripEnabled(False)

        # 顶层布局留白更充足，界面更像工作台而不是表单堆叠。
        self.verticalLayout_3.setContentsMargins(18, 18, 18, 18)
        self.verticalLayout_3.setSpacing(14)

        for layout in [
            self.verticalLayout_1,
            self.horizontalLayout_3,
            self.horizontalLayout_4,
            self.horizontalLayout_5,
            self.horizontalLayout_6,
            self.horizontalLayout_7,
            self.horizontalLayout_10,
            self.horizontalLayout,
        ]:
            layout.setSpacing(10)

        for label in [
            self.label_api_key,
            self.label_image_api_key,
            self.label_doc_list,
            self.label_doc_preview,
            self.label_design_method,
            self.label_industry,
            self.label_func_choice,
            self.label_param_choice,
            self.label_export,
            self.label_prompt,
            self.label_result,
        ]:
            label.setProperty("sectionLabel", True)

        for panel in [self.file_list, self.preview_area, self.prompt_input, self.result_area, self.plainTextEdit_update_talking]:
            panel.setProperty("panel", True)

        self.file_list.setProperty("densePanel", True)
        self.combo_kb.setProperty("primaryInput", True)
        self.comboBox_design_method.setProperty("primaryInput", True)
        self.label_stage.setProperty("statusBadge", True)
        self.label.setProperty("hintLabel", True)

        self.btn_add_kb.setProperty("secondaryButton", True)
        self.btn_refresh.setProperty("secondaryButton", True)
        self.btn_select_all.setProperty("secondaryButton", True)
        self.btn_clear_all.setProperty("secondaryButton", True)
        self.refresh_prompt_btn.setProperty("accentButton", True)
        self.pushButton_start_analyzer_image.setProperty("accentButton", True)
        self.export_btn.setProperty("successButton", True)
        self.pushButton_stop_generate.setProperty("dangerButton", True)

        self.api_key_input.setPlaceholderText("输入 DeepSeek API Key")
        self.lineEdit_image_api_key.setPlaceholderText("输入图片分析 API Key")
        self.combo_kb.setPlaceholderText("选择或切换需求目录")
        self.preview_area.setPlaceholderText("选择需求文件后，这里会显示提取的文档内容。")
        self.prompt_input.setPlaceholderText("这里会自动生成提示词，也可以手动微调。")
        self.result_area.setPlaceholderText("点击“开始推理”后，这里会展示生成结果。")
        self.plainTextEdit_update_talking.setPlaceholderText("这里会实时输出推理过程和日志。")

        self.file_list.setAlternatingRowColors(True)
        self.file_list.setFrameShape(QFrame.NoFrame)
        self.preview_area.setFrameShape(QFrame.NoFrame)
        self.prompt_input.setFrameShape(QFrame.NoFrame)
        self.result_area.setFrameShape(QFrame.NoFrame)
        self.plainTextEdit_update_talking.setFrameShape(QFrame.NoFrame)

        self._apply_panel_shadow(self.file_list)
        self._apply_panel_shadow(self.preview_area)
        self._apply_panel_shadow(self.prompt_input)
        self._apply_panel_shadow(self.result_area)
        self._apply_panel_shadow(self.plainTextEdit_update_talking)

        self.statusBar().showMessage("工作台已就绪，可先选择需求目录开始使用。", 5000)

    def _apply_panel_shadow(self, widget):
        effect = QGraphicsDropShadowEffect(self)
        effect.setBlurRadius(28)
        effect.setOffset(0, 8)
        effect.setColor(QtGui.QColor(15, 23, 42, 28))
        widget.setGraphicsEffect(effect)

    def _connect_signals(self):
        """连接所有信号和槽，集中管理，方便添加新按钮"""
        # 现有按钮连接
        self.btn_add_kb.clicked.connect(self.add_knowledge_base)
        self.btn_refresh.clicked.connect(self.load_saved_paths)
        self.combo_kb.currentTextChanged.connect(self.load_directory)
        self.file_list.itemSelectionChanged.connect(self.update_preview)
        self.btn_select_all.clicked.connect(lambda: self.file_list.selectAll())
        self.btn_clear_all.clicked.connect(lambda: self.file_list.clearSelection())
        self.generate_btn.clicked.connect(self.generate_report)
        self.refresh_prompt_btn.clicked.connect(self.generate_testcase_prompt)
        self.export_btn.clicked.connect(self.export_result)
        self.pushButton_stop_generate.clicked.connect(self.stop_generate)
        self.pushButton_clear_history_path.clicked.connect(self.clear_directory_to_combox)
        self.pushButton_start_analyzer_image.clicked.connect(self.start_image_analyzer)
        # 如果添加了新按钮，在这里添加连接
        # 示例：
        # if hasattr(self, 'new_button'):
        #     self.new_button.clicked.connect(self.new_button_handler)

    def updateLabel(self):
        self.job_area = self.comboBox.currentText()

    def generate_testcase_prompt(self, params=None):

        """
        生成测试用例设计提示词的智能函数
        生成测试用例设计提示词
        参数：
        params : dict/list - 参数维度字典或需求文档类型

        返回：
        str - 结构化提示词模板
        """
        # ========== 参数处理模块 ==========
        # 获取选择的方法，如果未选择方法则使用默认值
        methods = self.comboBox_design_method.currentText()
        if methods is None or methods == "":
            method = '常用测试用例设计方法'
            method_list = []
        else:
            method = methods
            method_list = method.split(',')

        parameters = ""
        func_type = self.func_choice_combo.currentText()
        if func_type == '接口测试用例':
            prompt = """# Role: 高级接口测试架构师

## Profile
- language: 中文
- description: 专业的接口质量保障专家，擅长设计高覆盖率的自动化测试方案
- background: 10年互联网企业测试架构经验，主导过百万级API测试体系建设
- personality: 严谨细致|逻辑性强|注重风险预防
- expertise: 测试策略设计|安全渗透测试|性能调优
- target_audience: 测试工程师|开发人员|DevOps团队

## Skills

1. 核心测试能力
   - 边界分析: 精准识别参数边界条件
   - 场景建模: 构建真实业务场景测试模型
   - 漏洞检测: OWASP TOP10安全漏洞测试
   - 性能压测: 设计阶梯式压力测试方案

2. 辅助技能
   - 协议解析: 深度理解HTTP/HTTPS/WebSocket协议
   - 数据构造: 生成结构化测试数据集
   - 异常模拟: 制造网络异常和服务端故障
   - 监控分析: 实时监测系统资源指标

## Rules

1. 质量原则：
   - 等价类划分必须覆盖有效/无效等价类
   - 边界值测试必须包含上点/离点/临界值
   - 安全测试需包含至少3种攻击向量
   - 性能测试需定义明确SLA指标

2. 行为准则：
   - 优先验证核心业务流程
   - 每个参数至少设计8个测试场景
   - 必须验证非功能需求
   - 保持测试用例原子性

3. 限制条件：
   - 不处理未文档化的接口
   - 不生成破坏性测试数据
   - 不泄露敏感测试信息
   - 不超出给定接口范围
   - 构造参数超长时不要使用*及repeat形式，直接展示需要多少字符即可，让用户自行修改
   - 严格按照示例输出，不要出现幻觉
   - 不要添加注释及多余的空行

## Workflows

- 目标: 生成零缺陷测试方案
- 步骤1: 解构接口文档要素
- 步骤2: 设计参数矩阵表
- 步骤3: 构建测试决策树
- 预期结果: 可执行的测试套件

## OutputFormat

1. 结构化输出：
   - format: application/json
   - structure: 只包含测试用例数组，不要添加注释
   - style: 符合Google JSON风格指南

2. 格式规范：
   - indentation: 2空格缩进
   - sections: 分测试类型分组
   - highlighting: 关键字段加注释

3. 验证规则：
   - validation: 通过ajv校验


4. 示例说明：
   1. 功能测试示例：
      - 标题: 用户信息查询验证
      - 格式类型: JSON
      - 示例内容:
        [{
          "case_id": "TC_FUNC_001",
          "case_name": "有效用户ID查询",
          "priority": "High",
          "pre_condition": "测试用户已注册并激活",
          "steps": [
            "构造合法用户ID请求",
            "验证200状态码",
            "校验响应数据结构",
            "比对数据库记录"
          ],
          "expected_result": "返回用户完整档案信息",
          "test_data": {
            "user_id": 10001,
            "auth_token": "Bearer valid_token"
          },
          "test_type": "功能测试"
        }]
   2. 安全测试示例：
      - 标题: SQL注入检测
      - 格式类型: JSON
      - 示例内容:
          [{
            "case_id": "TC_SEC_005",
            "case_name": "用户ID参数注入检测",
            "priority": "Critical",
            "pre_condition": "启用WAF防护",
            "steps": [
              "发送恶意注入载荷: ' OR 1=1 --",
              "监控数据库查询日志",
              "分析响应内容"
            ],
            "expected_result": "返回400错误且阻断注入请求",
            "test_data": {
              "user_id": "123' UNION SELECT * FROM users --"
            },
            "test_type": "安全测试"
          }]

## Initialization
作为高级接口测试架构师，你必须遵守上述Rules，按照Workflows执行任务，并按照[输出格式]输出。  """
            self.prompt_input.clear()
            self.prompt_input.setText(prompt)
        elif func_type == '功能测试用例':
            parameters += f"输出用例类型{func_type}"
            if isinstance(params, dict) and len(params) > 0:
                # 显式参数模式
                parameters = "参数维度：\n" + "\n".join(
                    [f"▸ {k}：{', '.join(v)}" for k, v in params.items()]
                )
            elif isinstance(params, list):
                # 需求文档类型提示
                doc_type = params[0] if params else "通用需求"
                parameters += f"需求文档类型：{doc_type}\n" + \
                              "请提取以下要素：\n" + \
                              "1. 核心业务实体及其属性\n" + \
                              "2. 关键业务流程步骤\n" + \
                              "3. 状态转换规则\n" + \
                              "4. 输入验证规则\n" + \
                              "5. 错误处理策略"
            else:
                # 默认需求分析模式
                parameters += "需求分析指南：\n" + \
                              "1. 识别功能边界（系统做什么/不做什么）\n" + \
                              "2. 提取业务规则（计算规则、验证规则）\n" + \
                              "3. 定义用户角色及其权限\n" + \
                              "4. 梳理关键业务流程（正常流、备选流、异常流）\n" + \
                              "5. 标记敏感操作（审计日志、权限校验点）"

            method_library = {
                "正交分析法": {
                    "desc": "使用正交表生成参数组合，覆盖所有参数对的交互组合",
                    "steps": ["构建正交表", "优化组合数量", "验证两两覆盖"],
                    "coverage": "参数对组合覆盖率 ≥95%"
                },
                "边界值分析": {
                    "desc": "针对数值型参数测试极值：最小值、略高于最小值、正常值、略低于最大值、最大值",
                    "steps": ["识别边界参数", "生成六点值（min-1,min,min+1,norm,max-1,max）", "处理无效类"],
                    "coverage": "边界条件覆盖率100%"
                },
                "等价类划分": {
                    "desc": "将输入划分为有效/无效类，每个类选取代表值测试",
                    "steps": ["定义有效等价类", "定义无效等价类", "生成代表值"],
                    "coverage": "每个等价类至少1个用例"
                },
                "状态转换": {
                    "desc": "基于状态机模型测试合法/非法转换",
                    "steps": ["绘制状态图", "覆盖所有合法转换", "测试非法转换"],
                    "coverage": "状态转换覆盖率100%"
                },
                "决策表": {
                    "desc": "条件组合的全覆盖测试（适合复杂业务规则）",
                    "steps": ["列出所有条件桩", "构建真值表", "合并相似项"],
                    "coverage": "条件组合覆盖率100%"
                },
                "错误推测": {
                    "desc": "基于经验测试易错点：异常输入、中断测试、并发操作",
                    "steps": ["列出历史缺陷", "分析脆弱模块", "设计非常规操作"],
                    "coverage": "补充覆盖边界外的5%"
                },
                "场景法": {
                    "desc": "模拟用户旅程测试端到端流程",
                    "steps": ["识别主成功场景", "定义扩展场景", "组合异常路径"],
                    "coverage": "主流程覆盖率100%"
                },
                "因果图": {
                    "desc": "分析输入条件的逻辑关系生成用例",
                    "steps": ["识别原因和结果", "构建因果图", "生成判定表"],
                    "coverage": "因果逻辑覆盖率100%"
                }
            }
            desc_str = ''
            # ========== 方法选择 ==========
            if len(method_list) >= 1:
                for method_item in method_list:
                    selected_method = method_library.get(method_item.strip(), method_library["正交分析法"])
                    desc_str += f"""
使用{method_item.strip()}方法设计用例时要符合：{selected_method['desc']}

关键步骤：
 {chr(10).join([f'{i + 1}. {step}' for i, step in enumerate(selected_method['steps'])])}
示例：
 {self.generate_example(method_item.strip())} \n

质量标准：
 - {selected_method['coverage']}
 - 正向场景用例占比70%
 - 异常场景用例占比20%
 - 边界场景用例占比10%
 \n
                    """
                # 格式化方法名称用于显示
                if len(method_list) == 1:
                    method_display = method_list[0].strip()
                else:
                    method_display = "、".join([m.strip() for m in method_list])
            else:
                method_display = method

            # ========== 生成提示词 ==========
            prompt = f"""


Role: 测试用例设计专家

Rules:
推理要求：
1. 【推理过程极度精简】每个推理步骤用一句话概括，使用箭头(→)或符号表示逻辑关系
2. 【输出完整内容】确保本分段的内容完整输出
3. 【控制token使用】避免重复和冗余表述

设计目标：\n
通过{method_display}实现：\n

用例数量：\n
根据需求的字数来决定测试用例的数量,尽可能多，越详细越好；
最低条数限制：规则1：大型功能200~300条,中型功能100~200条,小型功能50~100条，极小型功能10~20条，规则2：用例条数限制为需求字数的5%，按照规则1或规则2的最大值来决定；
最高条数限制：比最低条数浮动增加30%；
**不可遗漏任何需求（必须遵循的强制规则）
\n

用例设计需遵循：\n
{desc_str} \n

参数：\n
{parameters} \n

输出要求：
0. 当回复token即将达到上限时，不要截断输出，不满足完整用例结构的测试用例直接废弃该条，一定要确保回复的结构是完整的。
1. 格式：结构化JSON,必须严格遵守JSON语法规范,严格遵守示例中的格式，不擅自修改任何结构，不允许新增字段
2. 不要使用JavaScript语法（如.repeat()方法）
3. 对于需要重复字符的情况，请直接写出完整字符串，例如："aaaaaaa..."而不是"a".repeat(7)
4. 字符串长度限制测试时，请使用描述性文字如"256个字符的a"，而不是实际生成256个
5. 强制严格要求，只允许出现用例编号、用例标题、前置条件、操作步骤、预期结果这5个字段，且顺序是强制要求(不需要写优先级 字段 也不需要写 测试数据 字段）：
   - 用例编号：<模块缩写>-<3位序号>
   - 用例标题：<测试点> [正例/反例/设计方法]
   - 前置条件：初始化状态描述（不需要编号）
   - 操作步骤：带编号的明确步骤
   - 预期结果：可验证的断言（预期结果只有一个不需要编号，如果多个结果，就分成多条用例）
6.在最终回复的内容一开始先回复```json，回复完成后再多回复内容```，因为开始的```json和末尾的```我需要用来之后的数据格式处理。
7.最终回复的格式示例如下```json[完整的字典,完整的字典]```
生成步骤：
1. 参数建模 → 2. 场景分析 → 3. 用例生成 → 4. 交叉校验

"""
            self.prompt_input.clear()
            self.prompt_input.setText(prompt)
        # return prompt

    def generate_example(self, method):
        """生成方法对应的示例"""
        examples = {
            "正交分析法": """[
        {
            "用例编号": "PAY-001",
            "用例标题": "支付功能 [正例]",
            "前置条件": "用户已登录，购物车内已有商品",
            "操作步骤": [
                "1. 打开购物车页面",
                "2. 点击结算按钮",
                "3. 选择支付方式为支付宝支付",
                "4. 确认支付金额为100-1000元人民币",
                "5. 点击支付按钮"
            ],
            "预期结果": "支付成功，页面显示支付完成信息，余额扣减正确",
            "优先级": "P1"
        }
    ]""",
            "边界值分析": """[
        {
            "用例编号": "INPUT-001",
            "用例标题": "输入字段长度校验 [边界值测试]",
            "前置条件": "系统显示用户注册页面",
            "操作步骤": [
                "1. 打开注册页面",
                "2. 在用户名字段输入1个字符",
                "3. 在密码字段输入8个字符",
                "4. 在邮箱字段输入有效邮箱地址",
                "5. 点击提交按钮"
            ],
            "预期结果": "系统提示注册成功",
            "优先级": "P1"
        },
        {
            "用例编号": "INPUT-002",
            "用例标题": "输入字段长度校验 [反例，超长输入]",
            "前置条件": "系统显示用户注册页面",
            "操作步骤": [
                "1. 打开注册页面",
                "2. 在用户名字段输入超过50个字符",
                "3. 在密码字段输入8个字符",
                "4. 在邮箱字段输入有效邮箱地址",
                "5. 点击提交按钮"
            ],
            "预期结果": "系统提示用户名长度超限",
            "优先级": "P2"
        }
    ]""",
            "等价类划分": """[
        {
            "用例编号": "LOGIN-001",
            "用例标题": "登录功能 [有效等价类]",
            "前置条件": "用户已注册",
            "操作步骤": [
                "1. 打开登录页面",
                "2. 输入用户名为valid_user",
                "3. 输入密码为correct_password",
                "4. 点击登录按钮"
            ],
            "预期结果": "登录成功，跳转到首页",
            "优先级": "P1"
        },
        {
            "用例编号": "LOGIN-002",
            "用例标题": "登录功能 [无效等价类]",
            "前置条件": "用户已注册",
            "操作步骤": [
                "1. 打开登录页面",
                "2. 输入用户名为invalid_user",
                "3. 输入密码为random_password",
                "4. 点击登录按钮"
            ],
            "预期结果": "登录失败，提示用户名或密码错误",
            "优先级": "P1"
        }
    ]""",
            "状态转换": """[
        {
            "用例编号": "ORDER-001",
            "用例标题": "订单状态转换 [正常流程]",
            "前置条件": "用户购物车内有商品，订单创建成功",
            "操作步骤": [
                "1. 用户点击付款按钮",
                "2. 系统执行支付操作",
                "3. 支付成功后更新订单状态"
            ],
            "预期结果": "订单状态从'已创建'变为'已支付'",
            "优先级": "P1"
        },
        {
            "用例编号": "ORDER-002",
            "用例标题": "订单状态转换 [非法状态]",
            "前置条件": "订单状态为已取消",
            "操作步骤": [
                "1. 用户尝试付款已取消的订单",
                "2. 系统拦截付款请求"
            ],
            "预期结果": "操作失败，提示订单已取消，无法付款",
            "优先级": "P2"
        }
    ]""",
            "决策表": """[
        {
            "用例编号": "DISCOUNT-001",
            "用例标题": "会员折扣规则 [决策表测试]",
            "前置条件": "系统具有会员等级和折扣规则",
            "操作步骤": [
                "1. 用户登录账号，确认为黄金会员",
                "2. 添加商品到购物车，消费金额为500元",
                "3. 点击结算按钮"
            ],
            "预期结果": "系统计算折扣，实际支付金额为450元",
            "优先级": "P1"
        }
    ]""",
            "错误推测": """[
        {
            "用例编号": "UPLOAD-001",
            "用例标题": "文件上传 [异常输入]",
            "前置条件": "用户已登录，进入文件上传页面",
            "操作步骤": [
                "1. 用户选择一个exe文件",
                "2. 点击上传按钮"
            ],
            "预期结果": "系统提示不支持的文件类型，上传失败",
            "优先级": "P2"
        }
    ]""",
            "场景法": """[
        {
            "用例编号": "CHECKOUT-001",
            "用例标题": "用户购买商品 [主成功场景]",
            "前置条件": "用户已登录，购物车内有商品",
            "操作步骤": [
                "1. 用户进入购物车页面",
                "2. 点击结算按钮",
                "3. 填写收货地址",
                "4. 选择支付方式为支付宝",
                "5. 确认订单并付款"
            ],
            "预期结果": "订单支付成功，显示订单详情",
            "优先级": "P1"
        }
    ]""",
            "因果图": """[
        {
            "用例编号": "LOGIN-003",
            "用例标题": "登录功能 [因果关系测试]",
            "前置条件": "系统有登录模块",
            "操作步骤": [
                "1. 用户输入用户名为admin",
                "2. 输入密码为correct_password",
                "3. 点击登录按钮"
            ],
            "预期结果": "登录成功，跳转到管理页面",
            "优先级": "P1"
        },
        {
            "用例编号": "LOGIN-004",
            "用例标题": "登录功能 [因果关系测试 - 异常输入]",
            "前置条件": "系统有登录模块",
            "操作步骤": [
                "1. 用户输入用户名为admin",
                "2. 输入密码为wrong_password",
                "3. 点击登录按钮"
            ],
            "预期结果": "登录失败，提示用户名或密码错误",
            "优先级": "P1"
        }
    ]"""
        }
        return examples.get(method, "此方法示例未实现")

    def add_knowledge_base(self):
        """ 添加知识库目录（修改版）"""
        directory = QFileDialog.getExistingDirectory(self, "选择知识库目录")
        if directory:
            if directory not in self.get_current_knowledge_paths():
                self.add_directory_to_combox(directory)
                self.save_paths_to_config()  # 新增持久化存储

    def load_directory(self, directory):
        """ 加载目录文件（修正版）"""
        try:
            self.current_dir = directory
            self.file_list.clear()

            if not directory:
                return

            if not os.path.exists(directory):
                QMessageBox.warning(self, "路径错误", f"目录不存在: {directory}")
                return

            # 优化文件过滤逻辑
            # valid_extensions = ('.docx', '.xlsx', '.md', '.txt', '.pdf', '.json', 'yml', 'yaml')
            valid_extensions = ('.docx','.pdf')
            for fname in sorted(os.listdir(directory)):
                full_path = os.path.join(directory, fname)
                if os.path.isfile(full_path) and fname.lower().endswith(valid_extensions):
                    item = QListWidgetItem(fname)
                    item.setData(Qt.UserRole, full_path)
                    item.setToolTip(full_path)  # 添加路径提示
                    self.file_list.addItem(item)

            # 添加数量提示
            self.statusBar().showMessage(f"已加载 {self.file_list.count()} 个文档", 3000)
        except Exception as e:
            QMessageBox.critical(self, "加载错误", str(e))

    def update_preview(self):
        """ 更新预览内容 """
        self.preview_area.clear()
        self.selected = [item.data(Qt.UserRole) for item in self.file_list.selectedItems()]
        if not self.selected:
            return
        path = self.selected[0]
        # 检查文件是否为临时文件，否则将引起闪退
        if os.path.basename(path).startswith('~$'):
            QMessageBox.warning(self, "提示", "先关闭该文档，刷新需求列表后重新选择。")
            return

        # 命中缓存则直接展示
        cached = self._context_cache.get(path)
        if cached is not None:
            self.context = cached
            self.preview_area.setText(self.context)
            return

        # 注释掉下方的原有实现方案，直接复用现有方法
        if path.endswith('docx'):
            # 复用替换word图片的方法来获取纯文本
            self.context = insert_image_position_with_list(path, [])
        elif path.endswith('pdf'):
            self.context = extract_pdf_text_with_image_list(path, [])
        else:
            self.context = ""

        self._context_cache[path] = self.context
        self.preview_area.setText(self.context)

        # for path in self.selected:
        #     raw_text = self.read_file(path)
        #     # 根据文件类型处理内容
        #     if path.endswith('docx'):
        #         if not self.module_input.text():  # 未获取到文本标签，则使用默认清洗规则
        #             # cleaned = self.clean_text(raw_text)  # 注释掉未实现的方法
        #             content.append(raw_text)
        #         else:
        #             content.append({"paragraphs": raw_text})  # 拼接数据
        #     elif path.endswith('.pdf'):
        #         # PDF文件直接返回文本内容，需要特殊处理
        #         content.append(raw_text)  # PDF读取函数已返回正确格式
        #     else:
        #         content.append(raw_text)
        #
        # # 根据文件类型更新预览内容
        # for path in self.selected:
        #     if path.endswith('docx'):
        #         if not self.module_input.text() and path.endswith('docx'):  # 不指定，使用默认配置进行清洗
        #             for ele in content:
        #                 print("docx文档ele:", ele)
        #                 if isinstance(ele, dict) and "paragraphs" in ele:
        #                     print("docx文档ele[paragraphs]:", ele["paragraphs"])
        #                     all_content += "\n".join(ele["paragraphs"])
        #                 else:
        #                     all_content += str(ele)
        #             self.context = chunk_text(all_content)  # 不输入文本标题时，也对文本进行分块
        #         elif self.module_input.text() and path.endswith('docx'):  # 指定标题获取文档内容
        #             try:
        #                 for ele in content:
        #                     if isinstance(ele, dict) and "paragraphs" in ele:
        #                         paragraph = ele['paragraphs']
        #                         for key, value in paragraph.items():
        #                             all_content += f'{str(value)} \n'  # 获取指定标题内容
        #                 self.context = chunk_text(all_content)
        #             except Exception as e:
        #                 QMessageBox.critical(self, "预览", f"更新预览内容失败，错误信息{e}！")
        #     elif path.endswith('.pdf'):
        #         # 处理PDF文件内容
        #         for ele in content:
        #             if isinstance(ele, dict) and "paragraphs" in ele:
        #                 # 处理PDF返回的结构化数据
        #                 for paragraph in ele["paragraphs"]:
        #                     all_content += str(paragraph) + "\n"
        #             else:
        #                 all_content += str(ele)
        #         self.context = chunk_text(all_content)
        #     elif path.split('.')[-1].lower() in ('txt', 'xlsx', 'md'):
        #         if isinstance(content, list) and len(content) > 0:
        #             if isinstance(content[0], dict) and "paragraphs" in content[0]:
        #                 all_content = "\n".join(content[0]["paragraphs"])
        #             else:
        #                 all_content = content[0] if isinstance(content[0], str) else str(content[0])
        #     elif path.split('.')[-1].lower() in ('json', 'yml', 'yaml'):
        #         if isinstance(content, list) and len(content) > 0:
        #             all_content = content[0].get('paragraphs', '获取内容失败') if isinstance(content[0], dict) else str(
        #                 content[0])
        #             self.context = self.chunk_json(all_content) if isinstance(all_content,
        #                                                                       (dict, list)) else chunk_text(
        #                 str(all_content))
        #             all_content = json.dumps(all_content, indent=4, ensure_ascii=False) if isinstance(all_content, (
        #                 dict, list)) else str(all_content)
        #

    def read_file(self, file_path):
        """ 多格式文件读取 """
        try:
            if file_path.endswith('.docx'):
                try:
                    if not self.module_input.text():  # 如果没有文本标题输入
                        # 需要实现 extract_content 方法
                        doc = self.extract_content(file_path)
                    else:
                        # 需要实现 extract_text_by_title 方法
                        doc = self.extract_text_by_title(file_path,
                                                         self.module_input.text(),  # 文本标题
                                                         self.module_input_table.text(),  # 表格标题
                                                         self.module_input_pic.text())  # 图片标题
                    return doc
                except Exception as e:
                    print(f"读取docx文件出错: {e}")
                    import traceback
                    traceback.print_exc()
                    return ""
            elif file_path.endswith('.pdf'):
                try:
                    from PyPDF2 import PdfReader
                    with open(file_path, 'rb') as f:
                        reader = PdfReader(f)
                        pages_text = []
                        for page in reader.pages:
                            try:
                                # 提取页面文本
                                text = page.extract_text()
                                # 处理可能的编码问题
                                if text:
                                    # 移除特殊控制字符
                                    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
                                    # 处理常见的乱码字符
                                    text = text.encode('utf-8', errors='ignore').decode('utf-8')
                                pages_text.append(text if text else "")
                            except Exception as page_e:
                                print(f"提取PDF页面文本时出错: {page_e}")
                                pages_text.append("")
                        # 合并所有页面文本
                        combined_text = '\n'.join(pages_text)
                        # 清理多余的空白行
                        combined_text = '\n'.join(line for line in combined_text.split('\n') if line.strip())
                        # 返回与Word文档一致的结构化格式
                        return {"paragraphs": [combined_text], "tables": []}
                except Exception as e:
                    print(f"读取pdf文件出错: {e}")
                    import traceback
                    traceback.print_exc()
                    return {"paragraphs": [""], "tables": []}  # 保证返回一致的数据结构

            elif file_path.endswith('.xlsx'):
                try:
                    import pandas as pd
                    df = pd.read_excel(file_path)
                    # 将 DataFrame 转换为 Markdown 格式
                    excel_content = df.to_markdown(index=False)
                    return excel_content
                except Exception as e:
                    error_msg = f"无法读取Excel文件: {str(e)}"
                    print(error_msg)
                    import traceback
                    traceback.print_exc()
                    QMessageBox.critical(self, "错误", error_msg)
                    return ""
            elif file_path.endswith('.md'):
                try:
                    import markdown
                    with open(file_path, 'r', encoding='utf-8') as f:
                        return markdown.markdown(f.read())
                except Exception as e:
                    print(f"读取markdown文件出错: {e}")
                    import traceback
                    traceback.print_exc()
                    return ""
            elif file_path.endswith('.txt'):
                try:
                    # 尝试多种编码方式读取
                    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                content = f.read()
                            return content
                        except UnicodeDecodeError:
                            continue
                    # 如果所有编码都失败
                    raise Exception("无法使用常见编码读取文件")
                except Exception as e:
                    print(f"读取txt文件出错: {e}")
                    import traceback
                    traceback.print_exc()
                    return ""
            elif file_path.endswith('.json'):
                try:
                    import json
                    # 尝试多种编码方式读取
                    encodings = ['utf-8', 'gbk', 'gb2312']
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                return json.load(f)
                        except UnicodeDecodeError:
                            continue
                    # 如果所有编码都失败
                    raise Exception("无法使用常见编码读取JSON文件")
                except Exception as e:
                    print(f"读取json文件出错: {e}")
                    import traceback
                    traceback.print_exc()
                    return ""
            elif file_path.endswith(('.yaml', '.yml')):
                try:
                    import yaml
                    # 尝试多种编码方式读取
                    encodings = ['utf-8', 'gbk', 'gb2312']
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                return yaml.safe_load(f)
                        except UnicodeDecodeError:
                            continue
                    # 如果所有编码都失败
                    raise Exception("无法使用常见编码读取YAML文件")
                except Exception as e:
                    print(f"读取yaml文件出错: {e}")
                    import traceback
                    traceback.print_exc()
                    return ""
        except Exception as e:
            error_msg = f"读取文件时发生未知错误: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            QMessageBox.warning(self, "读取错误", error_msg)
            return ""

    def extract_content(self, file_path):
        """
        提取Word文档内容
        """
        try:
            from docx import Document
            doc = Document(file_path)
            content = {"paragraphs": [], "tables": []}

            # 提取段落内容
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content["paragraphs"].append(text)

            # 提取表格内容
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                content["tables"].append(table_data)

            return content
        except Exception as e:
            print(f"提取Word文档内容时出错: {e}")
            import traceback
            traceback.print_exc()
            return {"paragraphs": [], "tables": []}

    def extract_text_by_title(self, docx_path, title_keywords, table_keywords, pic_keywords):
        """
        根据标题提取Word文档内容
        """
        try:
            from docx import Document
            doc = Document(docx_path)
            result = {}

            for title_keyword in title_keywords.split(','):
                content = []
                capture = False
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        # 判断是否是标题
                        if title_keyword in text:
                            capture = True
                            content.append(text)
                            continue
                        # 捕获正文内容
                        if capture and text:
                            content.append(text)
                result[title_keyword] = "\n".join(content)
            return result
        except Exception as e:
            print(f"根据标题提取内容时出错: {e}")
            import traceback
            traceback.print_exc()
            return {}

    def generate_report(self):
        if self.generateButton.text() != "开始推理":
            return
        self.plainTextEdit_update_talking.clear()  # 清空上一次的AI回复对话框展示内容
        """ 生成分析报告 """
        if not self.prompt_input.toPlainText().strip():
            QMessageBox.warning(self, "提示", "请输入提示词！")
            return
        context = self.preview_area.toPlainText()
        if not context:
            QMessageBox.warning(self, "提示", "内容预览框中无数据，请求大模型终止！")
            return
        self.final_context = chunk_text(context)

        # 获取API Key
        self.api_key = self.api_key_input.text().strip()
        if not self.api_key:
            reply = QMessageBox.question(self, "提示", "未输入deepseek API Key，将使用默认配置。是否继续？",
                                         QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.No:
                return

        # 禁用按钮防止重复点击
        self.generate_btn.setEnabled(False)
        QApplication.setOverrideCursor(Qt.WaitCursor)
        self.job_area = self.comboBox.currentText()
        self.func_type = self.func_choice_combo.currentText()

        # 创建异步线程，传递API Key
        self.thread = GenerateThread(
            prompt=self.prompt_input.toPlainText(),
            context = self.final_context,
            job_area=self.job_area,
            func_type=self.func_type,
            design_method=self.comboBox_design_method.currentText(),
            api_key=self.api_key  # 传递API Key
        )
        self.thread.current_status.connect(self.update_talking)
        self.thread.finished.connect(self.on_generation_finished)
        self.thread.error.connect(self.on_generation_error)
        self.thread.current_stage.connect(self.update_generate_stage)

        try:
            if (not self.thread.isRunning()) and self.generateButton.text() == "开始推理":
                self.generateButton.setText("推理中...")
                self.thread.start()
        except Exception as e:
            # 启动失败时恢复 UI 状态，避免按钮卡死/鼠标转圈不恢复
            QMessageBox.critical(self, "启动失败", f"创建推理任务失败:\n{e}")
            self.generate_btn.setEnabled(True)
            self.generateButton.setText("开始推理")
            self.label_stage.setText("无进行中的推理")
            QApplication.restoreOverrideCursor()

    def on_generation_finished(self, result):
        """ 根据AI回答的内容，显示到结果组件中，修改相关组件状态 """
        # 确保结果显示为字符串
        if isinstance(result, (dict, list)):
            self.result_area.setText(json.dumps(result, indent=2, ensure_ascii=False))
        else:
            self.result_area.setText(str(result))
        self.generate_btn.setEnabled(True)
        self.generateButton.setText("开始推理")
        self.label_stage.setText("无进行中的推理")
        QApplication.restoreOverrideCursor()

    def on_generation_error(self, error_msg):
        """ 错误处理 """
        QMessageBox.critical(self, "生成错误", f"模型调用失败:\n{error_msg}")
        self.generate_btn.setEnabled(True)
        self.pushButton_start_analyzer_image.setEnabled(True)
        self.generateButton.setText("开始推理")
        self.label_stage.setText("无进行中的推理")
        QApplication.restoreOverrideCursor()

    def stop_generate(self):
        print("结束进程")
        try:
            if self.thread.isRunning():
                self.thread.terminate()  # 强制终止
                self.thread.wait()
                self.generate_btn.setEnabled(True)
        except AttributeError as e:
            print("线程未创建")
        finally:
            if self.generateButton.text() == "推理中...":
                self.generateButton.setText("开始推理")
            self.label_stage.setText("无进行中的推理")
            QApplication.restoreOverrideCursor()  # 停止鼠标转圈

    def start_image_analyzer(self):
        try:
            self.plainTextEdit_update_talking.clear()
            self.analyzer_enable = self.checkBox_analyzer_enable.isChecked()
            self.selected = [item.data(Qt.UserRole) for item in self.file_list.selectedItems()]
            # 获取image API Key
            self.image_api_key = self.lineEdit_image_api_key.text().strip()
            if not self.analyzer_enable or not self.selected:
                return
            if not self.image_api_key:
                reply = QMessageBox.question(self, "提示", "未输入通义千问 API Key，将使用默认配置。是否继续？",
                                             QMessageBox.Yes | QMessageBox.No)
                if reply == QMessageBox.No:
                    return
                elif reply == QMessageBox.Yes:
                    self.lineEdit_image_api_key.setText(config.constants.IMAGE_API_KEY)

            # 禁用ai分析按钮，避免重复点击
            self.pushButton_start_analyzer_image.setEnabled(False)
            pdf_path = self.selected[0]
            self._last_analyzed_path = pdf_path

            # 图片分析期间禁用“开始推理”，避免两套异步任务抢占同一个预览/状态
            self._generate_was_enabled = self.generate_btn.isEnabled()
            self.generate_btn.setEnabled(False)

            self.image_thread = ImageAnalyzer(pdf_path,batch_delay=1.0,image_api_key=self.image_api_key,analyzer_enable=self.analyzer_enable)
            self.image_thread.current_status.connect(self.update_talking)
            self.image_thread.finished.connect(self.on_analyzer_finished)
            self.image_thread.error.connect(self.on_generation_error)

            if not self.image_thread.isRunning():
                self.image_thread.start()
                print("图片分析线程已启动")
                self.generateButton.setText("Ai分析图片中...") #图片分析之前 禁止使用deepseek推理
        except KeyboardInterrupt as e:
            self.plainTextEdit_update_talking.appendPlainText("报错KeyboardInterrupt可能是中断了操作，请重新点击Ai分析按钮")

    def on_analyzer_finished(self,data):
        self.preview_area.setText(data) # 更新预览内容，把图片分析后
        self.context = data
        if self._last_analyzed_path:
            # 把分析后的结果写回缓存，避免用户重新选回同一文件时被旧预览覆盖
            self._context_cache[self._last_analyzed_path] = data
        self.pushButton_start_analyzer_image.setEnabled(True)
        self.generate_btn.setEnabled(self._generate_was_enabled)
        self.generateButton.setText("开始推理")
        # 分析完毕给，给一个弹框提示去内容预览里检查图片分析是否正
        QMessageBox.information(self,"ai分析完成","检查分析结果是否正确后，再点击开始推理")

    def _try_parse_json_from_text(self, text):
        """
        兼容 AI 输出文本的 JSON 提取：
        - 允许带 ```json ... ``` 代码块
        - 允许直接是纯 JSON
        - 允许 JSON 前后夹杂少量说明文字（尽量切片解析）
        """
        if not isinstance(text, str):
            return text

        import re
        import json

        raw = text.strip()
        if not raw:
            return None

        candidates = []
        # 优先提取 code block 内容
        codeblock_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", raw, flags=re.IGNORECASE)
        if codeblock_match:
            candidates.append(codeblock_match.group(1).strip())
        # 其次尝试直接解析整段
        candidates.append(raw)

        # 再次兜底：从第一个 '{' 或 '[' 切到最后一个对应括号
        start_obj = raw.find("{")
        start_arr = raw.find("[")
        start_candidates = [i for i in [start_obj, start_arr] if i != -1]
        if start_candidates:
            start = min(start_candidates)
            if raw[start] == "{":
                end = raw.rfind("}")
            else:
                end = raw.rfind("]")
            if end != -1 and end > start:
                candidates.append(raw[start:end + 1].strip())

        for candidate in candidates:
            try:
                return json.loads(candidate)
            except Exception:
                continue

        return None

    def json_to_excel(self, json_data, output_file):
        """
        将JSON数据转换为Excel文件，操作步骤中每个分号后添加换行符，并自动调整列宽和行高
        :param json_data: JSON数据（字符串或字典）
        :param output_file: 输出的Excel文件路径
        """
        try:
            import pandas as pd
            import json
            import math
            from openpyxl.utils import get_column_letter
            from openpyxl.styles import Alignment
            # 如果输入是JSON字符串，解析为字典
            if isinstance(json_data, str):
                data = self._try_parse_json_from_text(json_data)
                if data is None:
                    raise ValueError("导出XLSX失败：结果内容无法解析为有效JSON（可能缺少/多余格式包装）。")
            else:
                data = json_data

            # 处理多层嵌套的源数据
            data = normalize_data(data)

            # 如果数据是列表，直接转换为DataFrame
            if isinstance(data, list):
                # 格式化操作步骤，每个分号后添加换行符
                for i in data:
                    if isinstance(i, dict) and "操作步骤" in i:
                        if isinstance(i["操作步骤"], list):
                            # 如果是列表，每项后添加换行符
                            step = "\n".join([str(x) for x in i["操作步骤"]])
                            i["操作步骤"] = step
                        elif isinstance(i["操作步骤"], str):
                            # 如果是字符串，替换分号为换行符
                            i["操作步骤"] = i["操作步骤"].replace("; ", ";\n")

                df = pd.DataFrame(data)
            # 如果数据是字典，尝试提取其中的列表
            elif isinstance(data, dict):
                # 查找字典中的列表值
                list_data = None
                for key, value in data.items():
                    if isinstance(value, list):
                        list_data = value
                        # 格式化操作步骤，每个分号后添加换行符
                        for i in list_data:
                            if isinstance(i, dict) and "操作步骤" in i:
                                if isinstance(i["操作步骤"], list):
                                    step = "\n".join([str(x) for x in i["操作步骤"]])
                                    i["操作步骤"] = step
                                elif isinstance(i["操作步骤"], str):
                                    # 如果是字符串，替换分号为换行符
                                    i["操作步骤"] = i["操作步骤"].replace("; ", ";\n")
                        break
                if list_data:
                    df = pd.DataFrame(list_data)
                else:
                    # 如果没有找到列表，将字典转换为单行DataFrame
                    df = pd.DataFrame([data])
            else:
                raise ValueError("不支持的数据格式")

            # 保存为Excel文件并自动调整列宽和行高
            with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='TestCases')

                # 获取worksheet对象以调整列宽和行高
                worksheet = writer.sheets['TestCases']

                # 1) 自动调整列宽（放宽上限，避免中文长句被挤）
                col_widths = {}  # column letter -> width
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = get_column_letter(column[0].column)
                    for cell in column:
                        if cell.value is None:
                            continue
                        try:
                            cell_length = len(str(cell.value))
                            if cell_length > max_length:
                                max_length = cell_length
                        except Exception:
                            continue
                    # Excel 列宽单位近似为“字符数”，给一些 padding，适度上限
                    adjusted_width = min(max_length + 2, 120)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
                    col_widths[column_letter] = adjusted_width

                # 2) 强制换行 + 更稳的行高估算（同时考虑：显式换行 与 自动换行）
                #    Excel 的行高上限约 409，这里给到最大可用值，尽量保证不截断。
                MAX_EXCEL_ROW_HEIGHT = 409
                DEFAULT_ROW_HEIGHT = 15
                LINE_HEIGHT = 15

                for row in worksheet.iter_rows():
                    max_lines = 1
                    row_number = row[0].row

                    for cell in row:
                        if cell.value is None:
                            continue

                        text = str(cell.value)

                        # 显式换行优先：如果内容里已有 '\n'，就按其行数计算
                        if '\n' in text:
                            lines = text.count('\n') + 1
                        else:
                            # 估算自动换行行数：字符数 / 每行可容纳字符数
                            # 列宽=120左右时，每行大概可容纳的字符数更接近“列宽本身”，这里做个系数微调
                            col_letter = get_column_letter(cell.column)
                            width = col_widths.get(col_letter, 30)
                            chars_per_line = max(int(width * 1.1), 10)
                            lines = int(math.ceil(len(text) / chars_per_line))

                        max_lines = max(max_lines, lines)

                        # 强制 wrap，让 Excel 真正按估算换行展示
                        cell.alignment = Alignment(wrap_text=True, vertical='top')

                    worksheet.row_dimensions[row_number].height = min(
                        MAX_EXCEL_ROW_HEIGHT, max(DEFAULT_ROW_HEIGHT, max_lines * LINE_HEIGHT)
                    )

        except Exception as e:
            print(f"转换为Excel时出错: {e}")
            import traceback
            traceback.print_exc()

    def export_result(self):
        """ 导出结果 """
        if not self.result_area.toPlainText():
            QMessageBox.warning(self, "提示", "暂无导出内容，等待结果生成后导出")
            return

        # 基于“需求列表”选中的文件名生成默认导出文件名
        selected_paths = [item.data(Qt.UserRole) for item in self.file_list.selectedItems()]
        if selected_paths:
            first_path = selected_paths[0]
            base_name = os.path.splitext(os.path.basename(first_path))[0]
            if len(selected_paths) > 1:
                base_name = f"{base_name}_等{len(selected_paths)}个"
            default_dir = os.path.dirname(first_path)
        else:
            base_name = "导出结果"
            default_dir = os.getcwd()

        filter_str = self.get_export_filters()
        # 根据当前 filter 决定后缀
        ext_map = {
            "Excel Files (*.xlsx)": ".xlsx",
            "Word Documents (*.docx)": ".docx",
            "Text Files (*.txt)": ".txt",
            "Markdown Files (*.md)": ".md",
            "JSON Files (*.json)": ".json",
        }
        ext = ext_map.get(filter_str, ".txt")

        default_full_path = os.path.join(default_dir, f"{base_name}_导出{ext}")

        dlg = QFileDialog(self, "保存结果")
        dlg.setAcceptMode(QFileDialog.AcceptSave)
        dlg.setFileMode(QFileDialog.AnyFile)
        dlg.setNameFilter(filter_str)
        dlg.selectFile(default_full_path)
        if default_dir and os.path.exists(default_dir):
            dlg.setDirectory(default_dir)

        if dlg.exec_() != QDialog.Accepted:
            return

        path = dlg.selectedFiles()[0] if dlg.selectedFiles() else ""
        if not path:
            return

        try:
            content = self.result_area.toPlainText()
            real_ext = os.path.splitext(path)[1].lower()

            # JSON / XLSX 尝试从文本中提取有效 JSON（兼容 ```json 包装）
            parsed_obj = None
            if real_ext in [".json", ".xlsx"]:
                parsed_obj = self._try_parse_json_from_text(content)

            if real_ext == '.docx':
                from docx import Document
                doc = Document()
                # 如果是 JSON 结果，导出前做一次格式化，阅读体验更好
                if parsed_obj is not None:
                    doc.add_paragraph(json.dumps(parsed_obj, indent=2, ensure_ascii=False))
                else:
                    doc.add_paragraph(content)
                doc.save(path)
            elif real_ext == '.md':
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(content)
            elif real_ext == ".json":
                with open(path, 'w', encoding='utf-8') as file:
                    if parsed_obj is not None:
                        file.write(json.dumps(parsed_obj, indent=2, ensure_ascii=False))
                    else:
                        # 兜底：仍然写出原内容（同时提示用户可能不是合法 JSON）
                        file.write(content)
                        QMessageBox.warning(self, "警告", "结果内容未能解析为合法 JSON：已将原文本保存为 .json 文件。")
            elif real_ext == ".xlsx":
                if parsed_obj is None:
                    QMessageBox.warning(self, "导出失败", "结果内容无法解析为有效 JSON，无法导出 XLSX。")
                    return
                self.json_to_excel(parsed_obj, path)
            else:  # txt
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(content)

            QMessageBox.information(self, "导出成功", "文件已保存！")
        except Exception as e:
            QMessageBox.warning(self, "导出失败", str(e))

    def get_export_filters(self):
        """ 获取导出文件格式过滤器 """
        index = self.export_combo.currentIndex()
        return {
            0: "Excel Files (*.xlsx)",
            1: "Word Documents (*.docx)",
            2: "Text Files (*.txt)",
            3: "Markdown Files (*.md)",
            4: "JSON Files (*.json)"
        }[index]

    def chunk_json(self, content, max_chunk_size=1000):
        """
        将JSON内容按接口定义分块，确保接口数据完整。

        参数：
        - content: JSON内容（字典或列表）
        - max_chunk_size (int): 每个分块的最大字符数

        返回：
        - list: 分块后的接口定义列表
        """
        print("开始对JSON内容进行分块~~")
        chunks = []
        current_chunk = ""
        current_size = 0

        for n, interface in enumerate(content):
            interface_content = json.dumps(interface, indent=2, ensure_ascii=False)
            interface_size = len(interface_content)

            if current_size + interface_size > max_chunk_size:
                if n == 0:  # bug修复，当json对象列表的第一个元素尺寸就大于最大分隔尺寸时，则将第一个元素添加到分块列表中
                    chunks.append(interface_content)
                else:
                    chunks.append(current_chunk)
                current_chunk = interface_content
                current_size = interface_size
                n += 1
            else:
                current_chunk += "\n" + interface_content
                current_size += interface_size
                n += 1

        if current_chunk:
            chunks.append(current_chunk)

        return chunks

    def load_icon(self, icon_name):
        """
         智能加载图标，兼容开发环境和打包后环境
        :param icon_name:
        :return:
        """
        # 方法1: 打包后环境 - 资源在临时目录 sys._MEIPASS 中
        if hasattr(sys, '_MEIPASS'):
            base_path = sys._MEIPASS
            icon_path = os.path.join(base_path, 'config', icon_name)
            if os.path.exists(icon_path):
                return QtGui.QIcon(icon_path)

        # 方法2: 开发环境 - 尝试多个可能的相对路径
        possible_paths = [
            # 当前工作目录下的config文件夹
            os.path.join('config', icon_name),
            # 相对于脚本目录的config文件夹
            os.path.join(os.path.dirname(__file__), 'config', icon_name),
            # 您原来的路径（ui同级目录的config文件夹）
            os.path.join(os.path.dirname(__file__), '..', 'config', icon_name),
        ]

        for icon_path in possible_paths:
            if os.path.exists(icon_path):
                return QtGui.QIcon(icon_path)

        # 如果都找不到，输出错误信息（可选）
        print(f"警告: 未找到图标文件 {icon_name}")
        return QtGui.QIcon()  # 返回空图标

    def create_custom_model(self):
        """
        创建自定义多选下拉框
        :return:
        """
        for i in design_methods:
            item = QtGui.QStandardItem(i)
            item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
            if i == "因果图测试":
                item.setData(Qt.CheckState.Unchecked, Qt.ItemDataRole.CheckStateRole)
            else:
                # Qt.CheckState.checked 设置默认勾选的项
                item.setData(Qt.CheckState.Checked, Qt.ItemDataRole.CheckStateRole)
            self.custom_model.appendRow(item)
        self.comboBox_design_method.setModel(self.custom_model)
        self.comboBox_design_method.setCurrentIndex(-1)
        self.reset_combox_text()
        # 更新提示词
        self.generate_testcase_prompt()

    def reset_combox_text(self):
        """
        下拉框勾选有变化时，更新下拉框显示的内容
        :return:
        """
        # 创建一个列表，用于储存已勾选的项
        checked_list = []
        for i in range(self.custom_model.rowCount()):
            if self.custom_model.item(i).checkState() == Qt.CheckState.Checked:
                # 把项的文本内容 添加到用于储存已勾选的项的列表里
                checked_list.append(self.custom_model.item(i).text())
                # 把列表里的内容，显示在下拉框的文本框里，注意只能传入字符串，所以要把列表转化成字符串
                self.comboBox_design_method.setEditText(",".join(checked_list))
        self.generate_testcase_prompt()
        # 此处需要增加一个判断，当没有勾选任何项时，设置combox的索引为-1
        # 因为没有任何选项时，combox的文本框里会随便索引一个内容显示
        if not checked_list:  # 空列表在布尔表达式会被视为False，非空为True
            self.comboBox_design_method.setCurrentIndex(-1)

    def update_talking(self, data):
        """
        输出对话过程
        :param data:
        :return:
        """
        self.plainTextEdit_update_talking.appendPlainText(data)
        self.plainTextEdit_update_talking.moveCursor(QTextCursor.End)  # 滚动到底部

    def get_current_knowledge_paths(self):
        all_items_text = []  # 创建一个空列表来存储所有项的文本
        if self.combo_kb.count() > 0:
            for i in range(self.combo_kb.count()):
                item_text = self.combo_kb.itemText(i)
                all_items_text.append(item_text)
        return all_items_text

    def save_paths_to_config(self):
        """将当前选中的路径列表保存到配置"""
        # 获取所有知识库路径
        all_items_text = self.get_current_knowledge_paths()
        print("打印知识库目录:", all_items_text)
        # 使用QSettings保存列表
        self.settings.setValue("saved_directories", all_items_text)

    def load_saved_paths(self):
        """从配置加载之前保存的路径并添加到界面"""
        saved_paths = self.settings.value("saved_directories", [])
        print("读取saved_paths", saved_paths)
        # 注意：QSettings读取的列表可能是QStringList，确保转换为Python list of strings
        if isinstance(saved_paths, str):
            saved_paths = [saved_paths]  # 如果只有一个路径，它可能是字符串
        saved_paths = list(reversed(saved_paths))
        for path in saved_paths:
            self.add_directory_to_combox(path)  # 把目录添加到combox里
        self.load_directory(self.combo_kb.currentText())

    def add_directory_to_combox(self, path):
        """历史将知识库目录添加到combox中"""
        all_items_text = self.get_current_knowledge_paths()
        if path not in all_items_text:
            self.combo_kb.addItem(path)
            self.combo_kb.setCurrentText(path)

    def clear_directory_to_combox(self):
        # 清空setting里存储的saved_directories，设置为空列表就行
        self.settings.setValue("saved_directories", [])
        self.combo_kb.clear()  # 重新设置combox

    def update_generate_stage(self,stage):
        self.label_stage.setText(stage)